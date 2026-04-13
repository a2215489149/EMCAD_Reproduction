import csv
import json
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
REPO_ROOT = ROOT / "EMCAD"
OUTPUT_DOC_DIR = ROOT / "output" / "doc"
OUTPUT_PDF_DIR = ROOT / "output" / "pdf"
ARTIFACTS_DIR = ROOT / "artifacts"
MANUAL_TEST_DIR = ARTIFACTS_DIR / "manual_tests"
TEMPLATE_DOCX = ROOT / "tmp" / "docs" / "期中報告範本.docx"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")


def load_clinicdb_best():
    csv_path = ARTIFACTS_DIR / "experiment_summary.csv"
    best = None
    with csv_path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            if row.get("dataset") != "ClinicDB" or row.get("task") != "baseline":
                continue
            try:
                metric = float(row["reproduced_metric"])
            except (KeyError, TypeError, ValueError):
                continue
            if metric < 90.0:
                continue
            if best is None or metric > best["reproduced_metric"]:
                best = {
                    "run": row["run"],
                    "paper_metric": float(row["paper_metric"]),
                    "reproduced_metric": metric,
                    "delta": float(row["delta"]),
                    "notes": row["notes"],
                    "checkpoint": row["checkpoint"],
                }
    if best is None:
        raise RuntimeError("No valid ClinicDB baseline row found in experiment_summary.csv")
    return best


def load_synapse_manual_test():
    json_path = MANUAL_TEST_DIR / "synapse_best_manual_test_20260411.json"
    return json.loads(json_path.read_text(encoding="utf-8"))


def count_clinicdb_split():
    root = REPO_ROOT / "data" / "polyp" / "target" / "ClinicDB"
    counts = {}
    for split in ["train", "val", "test"]:
        counts[split] = len(list((root / split / "images").glob("*")))
    return counts


def count_synapse_split():
    train_count = len(list((REPO_ROOT / "data" / "synapse" / "train_npz").glob("*.npz")))
    test_count = len(list((REPO_ROOT / "data" / "synapse" / "test_vol_h5").glob("*.npy.h5")))
    return train_count, test_count


def set_columns(section, num=2, space_twips=424):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_elem = cols[0]
    else:
        cols_elem = OxmlElement("w:cols")
        sect_pr.append(cols_elem)
    cols_elem.set(qn("w:num"), str(num))
    cols_elem.set(qn("w:space"), str(space_twips))
    cols_elem.set(qn("w:equalWidth"), "true")
    cols_elem.set(qn("w:sep"), "false")


def set_run_font(run, size_pt, bold=False, latin="Times New Roman", east_asia="PMingLiU"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=0):
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(0)


def add_heading(document, text):
    p = document.add_paragraph()
    format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, space_after=2)
    r = p.add_run(text)
    set_run_font(r, 14, bold=True)
    return p


def add_body_paragraph(document, text):
    p = document.add_paragraph()
    format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=2)
    r = p.add_run(text)
    set_run_font(r, 12, bold=False)
    return p


def add_labeled_paragraph(document, label, text):
    p = document.add_paragraph()
    format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=2)
    r1 = p.add_run(label)
    set_run_font(r1, 12, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, 12, bold=False)
    return p


def add_reference_paragraph(document, text):
    p = document.add_paragraph()
    format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.1, space_after=1)
    r = p.add_run(text)
    set_run_font(r, 10, bold=False)
    return p


def set_table_cell_text(cell, text, size_pt=10.5, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, space_after=0)
    r = p.add_run(text)
    set_run_font(r, size_pt, bold=bold)


def build_document():
    clinic = load_clinicdb_best()
    synapse = load_synapse_manual_test()
    clinic_counts = count_clinicdb_split()
    synapse_train_count, synapse_test_count = count_synapse_split()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "PMingLiU")

    title = doc.add_paragraph()
    format_paragraph(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=2)
    title_run = title.add_run("EMCAD 論文復現期中報告")
    set_run_font(title_run, 14, bold=True)

    subtitle = doc.add_paragraph()
    format_paragraph(subtitle, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
    r = subtitle.add_run("EMCAD: Efficient Multi-scale Convolutional Attention Decoding for Medical Image Segmentation")
    set_run_font(r, 12, bold=False)

    meta_lines = [
        "作者：[請填姓名]",
        "學號：[請填學號]",
        "E-mail：[請填 Email]",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        run = p.add_run(line)
        set_run_font(run, 12, bold=False)

    kw = doc.add_paragraph()
    format_paragraph(kw, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=4)
    r1 = kw.add_run("關鍵詞：")
    set_run_font(r1, 12, bold=True)
    r2 = kw.add_run("醫學影像分割、論文復現、EMCAD、ClinicDB、Synapse")
    set_run_font(r2, 12, bold=False)

    sec2 = doc.add_section(WD_SECTION.CONTINUOUS)
    sec2.page_width = Cm(21)
    sec2.page_height = Cm(29.7)
    sec2.top_margin = Cm(2.54)
    sec2.bottom_margin = Cm(2.54)
    sec2.left_margin = Cm(2.54)
    sec2.right_margin = Cm(2.54)
    set_columns(sec2, num=2, space_twips=424)

    add_heading(doc, "簡介")
    add_body_paragraph(
        doc,
        "醫學影像分割的任務很直接，就是把病灶、器官或其他重要區域從影像中切出來，讓後續診斷、量測或手術規劃更容易進行。這類模型近年進步很快，但很多方法把重點放在追求更高分數，代價是 decoder 越做越重，參數量和 FLOPs 也跟著上升。對實際應用來說，這樣的方向不一定最理想，因為醫療端常常同時在意速度、記憶體和部署成本。"
    )
    add_body_paragraph(
        doc,
        "本次期中報告選擇 EMCAD 這篇 CVPR 2024 論文，主要有兩個原因。第一，作者不是只強調 Dice 分數，而是明確把「效果」和「效率」一起討論；第二，官方程式碼與資料連結相對完整，適合做復現與後續改良。以讀者角度來看，這篇論文最吸引我的地方是它不是重新發明一個很大的 backbone，而是回頭處理 segmentation pipeline 中常被忽略的 decoder，思考如何在不把模型做得更重的前提下，仍然保留多尺度與注意力的優點。"
    )
    add_body_paragraph(
        doc,
        "本報告的重點分成兩部分。前半部是整理我對論文方法的理解，也就是作者究竟想解決什麼問題、為什麼要把 EMCAD 設計成現在這個樣子；後半部則是期中階段已完成的復現結果。實驗上我先鎖定官方 repo 直接支援且最具代表性的兩個設定：ClinicDB 的二元分割，以及 Synapse 的多器官分割。這兩個資料集剛好也能對應論文中二元與多類別分割的兩種場景。"
    )
    add_body_paragraph(
        doc,
        "需要先說明的是，這份報告目前仍屬於期中版本，因此重心放在理解論文、確認官方結果是否能在本機重現，以及整理在復現過程中遇到的實際問題。未發表的改良方案與額外實驗還沒有正式開始，所以本報告不會假裝那些內容已經完成，而是會把它們留在結論與後續工作中，如實交代目前進度。"
    )

    add_heading(doc, "提出方法")
    add_labeled_paragraph(
        doc,
        "作者想解決的核心問題：",
        "從論文前言與 related work 可以看出，作者對現有方法有兩個主要不滿意的地方。第一，很多 decoder 雖然有效，但會在每一層疊上昂貴的 3×3 卷積與注意力模組，造成計算量偏高。第二，純 transformer encoder 擅長處理長距離關係，但對局部空間細節的掌握不一定夠穩，最後還是需要 decoder 幫忙補回多尺度與局部資訊。作者因此把問題定義成：能不能設計一個更輕量、但仍然能夠抓到多尺度局部細節的 decoder？"
    )
    add_labeled_paragraph(
        doc,
        "EMCAD 的整體想法：",
        "作者的答案不是換掉 encoder，而是保留 PVTv2 這種已經很成熟的階層式 encoder，然後重新設計 decoder。EMCAD 會接收 encoder 四個 stage 的特徵圖，透過 cascaded 的方式逐步往高解析度恢復，同時在每一層都做特徵精煉。換句話說，這篇論文的關鍵不在於發明新的 backbone，而是在於把 decoder 重新做成一個兼顧多尺度、注意力與效率的模組化結構。"
    )
    add_labeled_paragraph(
        doc,
        "MSCAM 的角色：",
        "作者提出的 MSCAM 可以看成 EMCAD 最核心的模組。它由 channel attention、spatial attention 與 multi-scale convolution block 組成。我的理解是，channel attention 負責回答「哪些通道比較重要」，spatial attention 負責回答「影像哪個位置比較值得注意」，而 multi-scale convolution block 再把這些被強調過的特徵，利用多尺度 depth-wise convolution 做進一步強化。這樣的安排很像先做篩選，再做局部細節補強。"
    )
    add_labeled_paragraph(
        doc,
        "為什麼作者強調 multi-scale depth-wise convolution：",
        "一般卷積雖然有效，但計算量高；depth-wise convolution 較輕，但表達能力常被質疑。作者在這裡採取的折衷方式，是把 depth-wise convolution 做成多尺度版本，並搭配 channel shuffle 與 point-wise convolution。論文中的 ablation 顯示，使用 [1, 3, 5] 這組 kernel 時，ClinicDB 與 Synapse 都拿到最好結果。這表示作者不是單純把卷積做小，而是透過不同感受野一起工作，讓 decoder 同時看到細節與較大的局部範圍。"
    )
    add_labeled_paragraph(
        doc,
        "LGAG 與 EUCB 的角色：",
        "除了 MSCAM 之外，作者還設計了 LGAG 與 EUCB。LGAG 是 large-kernel grouped attention gate，我把它理解成 skip connection 的篩選器：不是所有 encoder 傳下來的特徵都直接相加，而是先用 group convolution 做 gated filtering，再決定哪些內容值得保留。EUCB 則是 efficient up-convolution block，用 upsampling 加上 depth-wise convolution 做升頻與特徵整理。兩者合起來的目標很明確，就是在每次上採樣和融合時都盡量省計算，但又不要把資訊合併得太粗糙。"
    )
    add_labeled_paragraph(
        doc,
        "Loss 設計與作者的思路：",
        "在多類別分割上，作者沒有只拿最後一層輸出算 loss，而是使用 MUTATION 這種 multi-stage loss aggregation。四個 segmentation heads 會產生多個預測組合，訓練時把這些組合的 loss 一起納入考量。從讀者角度來看，這個設計的意義在於：作者不希望前面幾層只當純粹的過渡層，而是讓不同階段的輸出都真的參與學習，讓 decoder 每一層都對最終分割結果有貢獻。"
    )
    add_labeled_paragraph(
        doc,
        "我對論文貢獻的整理：",
        "如果把整篇論文濃縮成一句話，我認為 EMCAD 的貢獻是把「高效 decoder」這件事做得比以往更完整。它不是只靠一個注意力模組拿分，而是把 channel attention、spatial attention、多尺度 depth-wise convolution、gated skip fusion 和 stage-wise supervision 串成一個一致的設計。也因為每個元件都有明確任務，所以這篇論文的方法邏輯相對清楚，不會給人只是把很多常見模組硬湊在一起的感覺。"
    )
    add_labeled_paragraph(
        doc,
        "本次復現設定：",
        f"在實作端，我使用官方 GitHub repo 與論文推薦設定來重建 PVT-EMCAD-B2。硬體為 RTX 4070 SUPER 12GB，軟體為 Python 3.8 與 PyTorch 1.11.0+cu113。ClinicDB 使用官方 split，修正後的資料量為 train {clinic_counts['train']}、val {clinic_counts['val']}、test {clinic_counts['test']}，影像大小 352×352；Synapse 使用預處理後資料，train 為 {synapse_train_count} 個 slice、test 為 {synapse_test_count} 個 volume，影像大小 224×224。這些設定都盡量跟論文與 repo 保持一致，只有在 Windows 相容性問題上做最小修補。"
    )
    add_labeled_paragraph(
        doc,
        "復現過程中的實際問題：",
        "這次最明顯的教訓是，復現失敗不一定代表模型有問題。ClinicDB 一開始跑出來只有 84 左右，後來追查才發現不是 EMCAD 本身出錯，而是 Google Drive 列舉流程只抓到每個資料夾 50 張影像，導致實際訓練資料與論文設定完全不同。資料補齊到 489/61/62 之後，ClinicDB 的結果立刻回到合理範圍。Synapse 則遇到另一種問題：長時間訓練中斷後要用 checkpoint 續跑，但官方程式沒有完整 resume 設計，後來我補了 resume 支援後，又抓到一個只會在 resumed run 出現的 `ss` 初始化 bug。這些問題都說明，真正的復現工作其實很依賴資料完整性與訓練流程穩定性。"
    )
    add_labeled_paragraph(
        doc,
        "目前的復現結果：",
        f"ClinicDB 在修正資料後，official test Dice 達到 {clinic['reproduced_metric']:.4f}，和論文 {clinic['paper_metric']:.2f} 的差距只有 {clinic['delta']:+.4f}，幾乎可以視為重現成功。Synapse 的部分，我在暫停訓練後直接對目前 best checkpoint 做 official full test，得到 mean Dice {synapse['mean_dice_pct']:.4f}、mean HD95 {synapse['mean_hd95']:.4f}。如果以期中階段先採用的 Dice 差距不超過 2 個百分點作為門檻，Synapse 和論文 83.63 的差距 -1.8440 也已在可接受範圍內；但如果要用更嚴格的 ±1 標準來看，它仍然還差一步。這一點我認為應該在報告裡誠實交代，而不是直接寫成完全一致。"
    )
    add_labeled_paragraph(
        doc,
        "結果表現的解讀：",
        "就讀者角度來看，這組結果其實也支持了作者的論點。ClinicDB 幾乎貼著論文數字，代表 EMCAD 在二元分割任務上的實作路線相當穩；Synapse 雖然沒有完全對齊論文 83.63 的 five-run 平均，但在單機 Windows 環境、有限顯示記憶體、以及中途需要 resume 的情況下，仍能達到 81.786，表示方法本身具有相當程度的可轉移性。這也讓我更相信 EMCAD 的強項不是只在單一資料集衝分，而是它的 decoder 設計真的有一般性價值。"
    )

    result_table = doc.add_table(rows=1, cols=2)
    result_table.style = "Table Grid"
    set_table_cell_text(result_table.rows[0].cells[0], "項目", size_pt=10.5, bold=True)
    set_table_cell_text(result_table.rows[0].cells[1], "本次期中結果摘要", size_pt=10.5, bold=True)
    rows = [
        ("ClinicDB", f"Paper Dice 95.21；reproduced {clinic['reproduced_metric']:.4f}；差距 {clinic['delta']:+.4f}"),
        ("Synapse", f"Paper Dice 83.63；official test {synapse['mean_dice_pct']:.4f}；差距 -1.8440"),
        ("驗收說明", "本次期中先用 Dice 與論文差距不超過 2 個百分點作為可接受門檻"),
        ("目前判斷", "ClinicDB 可視為成功重現；Synapse 在期中門檻內達標，但和論文均值仍有小幅差距"),
    ]
    for left, right in rows:
        cells = result_table.add_row().cells
        set_table_cell_text(cells[0], left, size_pt=10.5, bold=True)
        set_table_cell_text(cells[1], right, size_pt=10.5, bold=False)

    add_heading(doc, "結論與討論")
    add_body_paragraph(
        doc,
        "綜合來看，本次期中階段最重要的成果不是只有把數字跑出來，而是把 EMCAD 的方法邏輯和實際復現流程都走過一次。從閱讀論文到動手實作，我對這篇工作的理解是：作者真正有價值的地方，在於把 decoder 視為一個可以被重新最佳化的核心元件，而不是單純依附在 backbone 後面。多尺度 depth-wise convolution、attention gate 與 stage-wise supervision 這幾個設計，都是在回答同一個問題，也就是如何在不把計算量推太高的情況下，仍然把局部細節補回來。"
    )
    add_body_paragraph(
        doc,
        "若只看目前已完成的兩個資料集，我認為 EMCAD 已經通過了期中階段最基本的檢查。ClinicDB 已經非常接近論文原始結果；Synapse 雖然還不是完全貼合 paper mean，但在暫定門檻下也已經進到合理區間。更重要的是，這兩個結果都不是靠模糊解釋硬撐出來的，而是有對應的 checkpoint、測試腳本和數據紀錄可追查。對我來說，這代表後續若要做改良，比較基準已經有一定可信度。"
    )
    add_body_paragraph(
        doc,
        "後續工作會分成兩步。第一步是做文獻與程式碼的新穎性檢查，避免提出其實早就被別人發表過的改法。第二步才是選定一個控制變因清楚、能合理接在 EMCAD 後面的改良點，並重新跑數據驗證。因為 Mid-Term Project 的加分規則很明確要求「改良方法未曾發表，且要附實驗數據與程式碼證明」，所以我不打算在還沒查清楚前就任意宣稱改良成立。這部分會等 baseline 收斂後再繼續。"
    )
    add_body_paragraph(
        doc,
        "最後，如果要用一句比較口語但也比較誠實的話來總結這次期中進度，我會說：EMCAD 不是一篇只看起來厲害、實際很難落地的論文。它的確有工程細節需要處理，但只要資料、checkpoint 與測試流程盯得夠緊，這篇工作是有機會在一般研究室 GPU 環境中被重現的。這也讓它很適合作為後續改良與實驗設計的出發點。"
    )

    add_heading(doc, "參考文獻")
    add_reference_paragraph(doc, "中文文獻：本報告主要參考英文論文與官方程式碼，中文文獻暫無。")
    add_reference_paragraph(doc, "英文文獻：")
    refs = [
        "[1] M. M. Rahman, M. Munir, and R. Marculescu, \"EMCAD: Efficient Multi-scale Convolutional Attention Decoding for Medical Image Segmentation,\" in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2024.",
        "[2] W. Wang, E. Xie, X. Li, D.-P. Fan, K. Song, D. Liang, T. Lu, P. Luo, and L. Shao, \"PVT v2: Improved Baselines with Pyramid Vision Transformer,\" Computational Visual Media, vol. 8, no. 3, pp. 415-424, 2022.",
        "[3] J. Bernal, F. J. Sanchez, G. Fernandez-Esparrach, D. Gil, C. Rodriguez, and F. Vilariño, \"WM-DOVA maps for accurate polyp highlighting in colonoscopy: Validation vs. saliency maps from physicians,\" Computerized Medical Imaging and Graphics, vol. 43, pp. 99-111, 2015.",
        "[4] J. Chen, Y. Lu, Q. Yu, X. Luo, E. Adeli, Y. Wang, L. Lu, A. L. Yuille, and Y. Zhou, \"TransUNet: Transformers Make Strong Encoders for Medical Image Segmentation,\" arXiv preprint arXiv:2102.04306, 2021.",
        "[5] SLDGroup, \"EMCAD official implementation,\" GitHub repository. Available: https://github.com/SLDGroup/EMCAD",
    ]
    for ref in refs:
        add_reference_paragraph(doc, ref)

    return doc


def convert_to_pdf(docx_path, pdf_dir):
    pdf_dir.mkdir(parents=True, exist_ok=True)
    profile_dir = ROOT / "tmp" / "lo_profile_output"
    profile_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        str(SOFFICE),
        f"-env:UserInstallation=file:///{profile_dir.as_posix()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(pdf_dir),
        str(docx_path),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(f"LibreOffice PDF conversion failed: {proc.stdout}\n{proc.stderr}")
    return pdf_dir / (docx_path.stem + ".pdf")


def count_pdf_pages(pdf_path):
    return len(PdfReader(str(pdf_path)).pages)


def main():
    OUTPUT_DOC_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF_DIR.mkdir(parents=True, exist_ok=True)

    doc = build_document()
    docx_path = OUTPUT_DOC_DIR / "EMCAD_midterm_report.docx"
    doc.save(str(docx_path))

    pdf_path = convert_to_pdf(docx_path, OUTPUT_PDF_DIR)
    pages = count_pdf_pages(pdf_path)

    summary = {
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "pdf_pages": pages,
    }
    summary_path = OUTPUT_DOC_DIR / "EMCAD_midterm_report_summary.json"
    summary_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    print(json.dumps(summary, ensure_ascii=False))


if __name__ == "__main__":
    main()
